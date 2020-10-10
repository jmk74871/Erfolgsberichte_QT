import csv

from docx import Document
from qtpy import QtCore


# ToDo:
#     find a way to use the filters from GUI


class ReportWriter:
    def __init__(self, from_date, to_date, user):
        self.user = user
        self.from_date = from_date
        self.to_date = to_date

    def create_report(self):
        report = Document()

        # creater base structure
        head = report.add_heading("Erfolgsbericht")
        report.add_paragraph('')
        report.add_paragraph(f'für: {self.user.name}')
        report.add_paragraph(f'Zeitraum: '
                             f'{self.from_date.day():02d}.{self.from_date.month():02d}.{self.from_date.year():04d} '
                             f'bis {self.to_date.day():02d}.{self.to_date.month():02d}.{self.to_date.year():04d}')
        report.add_paragraph('')

        ami = report.add_heading('Erfolge im Bereich Arbeitsmarktintegration', level=2)
        ea = report.add_heading('Erfolge im Bereich Ehrenamt', level=2)
        netz = report.add_heading('Erfolge im Bereich Vernetzung', level=2)
        other = report.add_heading('Sonstige Erfolge', level=2)
        end = report.add_paragraph(' ')

        # go through entrys and add them to report
        with open(self.user.db_path, mode='r', encoding='utf_8') as data:
            reader = csv.DictReader(data)

            for e in reader:
                q_date = QtCore.QDate(int(e['DATE'][0:4]), int(e['DATE'][5:7]), int(e['DATE'][8:]))
                if self.from_date <= q_date <= self.to_date:
                    if e['CATEGORY'] == "AMI" or e['CATEGORY'] == "Arbeitsmarktintegration":
                        ea.insert_paragraph_before(e['TEXT'], style='List Bullet')

                    elif e['CATEGORY'] == "EA" or e['CATEGORY'] == "Ehrenamt":
                        netz.insert_paragraph_before(e['TEXT'], style='List Bullet')

                    elif e['CATEGORY'] == "VN" or e['CATEGORY'] == "Vernetzung":
                        other.insert_paragraph_before(e['TEXT'], style='List Bullet')

                    else:
                        end.insert_paragraph_before(e['TEXT'], style='List Bullet')

        if self.from_date.year() == self.to_date.year() and self.from_date.month() == self.to_date.month():
            report.save(self.user.rep_dir + f'/{self.from_date.year():04d}-{self.from_date.month():02d}_'
                                            f'{self.user.name}.docx')
        elif self.from_date.year() == self.to_date.year() and self.from_date.month() != self.to_date.month():
            report.save(self.user.rep_dir + f'/{self.from_date.year():04d}-{self.from_date.month():02d}_to_'
                                            f'{self.to_date.month():02d}_{self.user.name}.docx')
        else:
            report.save(self.user.rep_dir + f'/{self.from_date.year():04d}-{self.from_date.month():02d}_to_'
                                            f'{self.to_date.year():04d}-{self.to_date.month():02d}_'
                                            f'{self.user.name}.docx')
